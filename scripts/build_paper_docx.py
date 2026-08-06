#!/usr/bin/env python3
"""
build_paper_docx.py - fill the ACM template with the drafted paper.

    python scripts/build_paper_docx.py

Works by EDITING a copy of `docs/DEEPLRN Group 2_ Milestone Proposal.docx`
rather than generating a document from scratch, so the ACM styling is
preserved exactly: Linux Biolinum, US Letter, 0.75in margins, the
single-column title block followed by a two-column body section, 11pt
section headings, 9pt keywords, 7pt references.

Placeholder paragraphs from the template are removed and replaced with the
drafted content and the generated figures. Nothing about the page setup,
fonts or column layout is changed.
"""

import copy
import os
import shutil

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

TEMPLATE = os.path.join("docs", "DEEPLRN Group 2_ Milestone Proposal.docx")
OUT = os.path.join("docs", "DEEPLRN Group 2_ Final Paper.docx")
FIG = os.path.join("results", "figures")

BODY_PT, HEAD_PT, SMALL_PT, REF_PT = 9.0, 11.0, 9.0, 7.0
FONT = "Linux Biolinum"


def clear_after(doc, keep_upto_text):
    """Delete every body paragraph after the one containing `keep_upto_text`."""
    keep = None
    for i, p in enumerate(doc.paragraphs):
        if keep_upto_text in p.text:
            keep = i
    for p in doc.paragraphs[keep + 1:]:
        p._element.getparent().remove(p._element)
    return doc.paragraphs[keep]


def para(doc, text="", size=BODY_PT, bold=False, italic=False, align=None,
         space_before=0, space_after=3):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if align is not None:
        p.alignment = align
    if text:
        add_runs(p, text, size, bold, italic)
    return p


def add_runs(p, text, size=BODY_PT, bold=False, italic=False):
    """Split on **bold** and *italic* markers."""
    import re
    for tok in re.split(r"(\*\*.+?\*\*|\*[^*]+?\*)", text):
        if not tok:
            continue
        if tok.startswith("**") and tok.endswith("**"):
            r = p.add_run(tok[2:-2]); r.bold = True
        elif tok.startswith("*") and tok.endswith("*"):
            r = p.add_run(tok[1:-1]); r.italic = True
        else:
            r = p.add_run(tok); r.bold = bold; r.italic = italic
        r.font.name = FONT
        r.font.size = Pt(size)
    return p


def heading(doc, text):
    """Numbered section heading, matching the template's 11pt normal style."""
    return para(doc, text, size=HEAD_PT, space_before=8, space_after=4)


def caps_heading(doc, text):
    return para(doc, text, size=HEAD_PT, bold=True, space_before=8, space_after=4)


def figure(doc, filename, caption, width_in=3.2):
    path = os.path.join(FIG, filename)
    if not os.path.isfile(path):
        print(f"  ! missing figure {path}")
        return
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(path, width=Inches(width_in))
    c = para(doc, caption, size=8.0, space_after=8)
    c.alignment = WD_ALIGN_PARAGRAPH.LEFT
    print(f"  + figure {filename}")


def _hrule(row, edges=("top", "bottom")):
    """Horizontal rule on a row. The template defines no table style, so
    borders are set directly rather than via a style that does not exist."""
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        borders = OxmlElement("w:tcBorders")
        for e in edges:
            el = OxmlElement(f"w:{e}")
            el.set(qn("w:val"), "single"); el.set(qn("w:sz"), "6")
            el.set(qn("w:color"), "000000")
            borders.append(el)
        tcPr.append(borders)


def table(doc, headers, rows, caption=None):
    t = doc.add_table(rows=1, cols=len(headers))
    for c, h in zip(t.rows[0].cells, headers):
        c.text = ""
        add_runs(c.paragraphs[0], h, size=7.5, bold=True)
    for row in rows:
        cells = t.add_row().cells
        for c, v in zip(cells, row):
            c.text = ""
            add_runs(c.paragraphs[0], str(v), size=7.5)
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                p.paragraph_format.space_before = Pt(0)
                p.paragraph_format.space_after = Pt(0)
    # booktabs-style rules: above and below the header, and under the last row
    _hrule(t.rows[0], ("top", "bottom"))
    _hrule(t.rows[-1], ("bottom",))
    if caption:
        para(doc, caption, size=8.0, space_before=2, space_after=8)


def bullets(doc, items, size=BODY_PT):
    """The template defines no list style (only 'normal' and headings), and its
    own CCS line uses a literal bullet, so we follow that convention with a
    hanging indent rather than introducing a style the template lacks."""
    for it in items:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Inches(0.16)
        p.paragraph_format.first_line_indent = Inches(-0.10)
        add_runs(p, "• " + it, size=size)


def main():
    if not os.path.isfile(TEMPLATE):
        raise SystemExit(f"template not found: {TEMPLATE}")
    shutil.copy2(TEMPLATE, OUT)
    doc = Document(OUT)

    # Keep the title block, replace the abstract, drop every placeholder after it.
    for p in doc.paragraphs:
        if p.text.strip().startswith("Unsupervised young children"):
            for r in p.runs[1:]:
                r._element.getparent().remove(r._element)
            p.runs[0].text = ABSTRACT
            break
    clear_after(doc, "Insert keyword text")

    # CCS and keywords still hold template placeholders; fill them.
    for p in doc.paragraphs:
        t = p.text.strip()
        if t.startswith("• Insert CCS"):
            for r in p.runs[1:]:
                r._element.getparent().remove(r._element)
            p.runs[0].text = CCS
        elif t.startswith("Insert keyword text"):
            for r in p.runs[1:]:
                r._element.getparent().remove(r._element)
            p.runs[0].text = KEYWORDS

    build_body(doc)
    doc.save(OUT)
    print(f"\nwrote {OUT}")


# ---------------------------------------------------------------- content --

ABSTRACT = (
    "Young children left briefly unattended can reach household hazards such as "
    "knives, tools or small swallowable objects. Existing home monitoring products "
    "detect movement into manually configured zones, but have no notion of which "
    "objects are dangerous or where those objects currently are. This work builds "
    "and evaluates a pipeline that detects children and hazardous objects using two "
    "separately trained nano-scale object detectors and combines their outputs with "
    "a geometric proximity rule, producing a Safe or Unsafe label for each frame. "
    "The two detectors reach 0.967 and 0.551 mean average precision at an "
    "intersection-over-union threshold of 0.5 on held-out test data, and run "
    "together in 7.2 milliseconds per frame, comfortably faster than real time. The "
    "proximity rule reaches a balanced accuracy of 0.595 on a held-out evaluation "
    "set, but its confidence interval includes chance, so the geometric stage is not "
    "confirmed as effective. Detection failure, not the choice of distance measure, "
    "is shown to be the dominant error source. Hyperparameter tuning selected on a "
    "shortened training schedule transferred poorly and reduced final accuracy."
)

CCS = ("• Computing methodologies → Object detection • Computing methodologies → "
       "Neural networks • Applied computing → Consumer health")

KEYWORDS = ("Object detection, YOLOv8, computer vision, child safety, "
            "proximity estimation, model fusion")


def build_body(doc):
    A = WD_ALIGN_PARAGRAPH

    # ---- 1 Introduction
    heading(doc, "1 Introduction")
    para(doc, "A child alone in a kitchen for a few minutes is enough time for a "
        "serious injury: a knife left within reach, a tool on a low shelf, or a coin "
        "small enough to swallow. This risk has driven a large consumer market in "
        "smart home cameras and baby monitors. Products in this category typically "
        "offer a “danger zone” feature, in which a caregiver manually marks a region "
        "of the home — near a staircase, a stove, or a pool — and receives an alert "
        "when the child enters it.")
    para(doc, "That design has a structural limitation. The zone is fixed at "
        "configuration time and the system reasons only about the child’s coordinates "
        "relative to it. It has no representation of what is dangerous, so a knife "
        "left on a low table outside the marked region produces no alert, while a "
        "child playing safely inside a marked region produces a false one. The hazard "
        "itself is invisible to the system. This work addresses that limitation by "
        "detecting the hazardous object as well as the child, and computing risk from "
        "the measured distance between them, with no manual configuration.")
    para(doc, "**Research question.** The pipeline needs two detection capabilities: "
        "locating children, and locating a diverse set of household hazards. These "
        "can be learned by one multi-class model or by two independently trained "
        "specialists. The latter avoids forcing a single network to share capacity "
        "between visually dissimilar tasks and lets each detector be configured "
        "independently, but requires running two models per frame. We ask whether "
        "**two independently trained specialist detectors, combined only at the output "
        "level by a geometric rule, form a workable child safety risk detector**, and "
        "at what computational cost.")
    para(doc, "An earlier framing asked whether this design *outperforms* a unified "
        "multi-class model. That comparison was abandoned because it cannot be "
        "answered with the available evidence: the published unified-model figure was "
        "obtained on a different dataset, with different hazard classes, using a "
        "substantially larger network. Establishing superiority would require training "
        "a unified model on the union of our own datasets as an internal control, "
        "which was outside the compute budget. The supportable claim is the weaker "
        "one: that the specialists perform *competitively* with published figures.")
    para(doc, "**Contributions.**")
    bullets(doc, [
        "A complete, reproducible two-detector pipeline with an output-level "
        "proximity rule, evaluated end to end on held-out data, including the cost "
        "of running both detectors per frame.",
        "An honest negative result on the risk-classification stage: its held-out "
        "confidence interval includes chance, and detection failure rather than "
        "distance measurement is identified as the dominant error source.",
        "Three methodological findings that generalise beyond this application — "
        "tuning on a shortened schedule can reduce final accuracy; a sequential "
        "ablation can invert when an earlier parameter is corrected; and a widely "
        "used framework default silently discards a hyperparameter under tuning.",
        "A measured account of dataset contamination in a public dataset, with its "
        "effect on reported accuracy quantified rather than assumed.",
    ])

    # ---- 2 Related Work
    heading(doc, "2 Related Work")
    para(doc, "[To be written. Requires at least eight peer-reviewed papers organised "
        "thematically; five are currently available. Themes to cover: automated "
        "child-safety monitoring; single-stage object detection and the accuracy/cost "
        "trade-off motivating nano-scale models; small-object detection, which is "
        "directly relevant since missed small hazards are our dominant failure mode; "
        "output-level model fusion versus joint multi-task training; proximity "
        "reasoning in safety-critical vision; and dataset contamination. Positioning: "
        "this work adopts the detect-then-measure-distance framework of prior "
        "child-safety systems but separates detection into two independently trained "
        "specialists and evaluates the geometric stage in isolation, which prior work "
        "has not done.]", italic=True)

    # ---- 3 Methodology
    heading(doc, "3 Methodology")
    para(doc, "**Architecture.** The pipeline has three stages, shown in Figure 1: "
        "two parallel detectors and a fusion rule with no learned parameters.")
    figure(doc, "architecture.png",
           "Figure 1: Pipeline architecture. The detectors are trained on separate "
           "datasets with no shared weights and no joint training, and use different "
           "input resolutions. Fusion is a threshold on normalised centroid distance.")
    para(doc, "The **child detector** is a YOLOv8-nano model fine-tuned on a single "
        "child class at 416×416. The **hazard detector** is a separate YOLOv8-nano "
        "model fine-tuned on twelve classes — Axe, Chainsaw, Chisel, Coin, Drink, "
        "Dumbbell, Fork, Hammer, Knife, Scissors, Screwdriver, Stapler — at 640×640. "
        "All twelve collapse to a single hazard category at fusion, but the specific "
        "class is retained for display, so the system reports “Knife” rather than "
        "“hazard”.")
    para(doc, "**Risk-fusion layer.** If either detector returns no boxes the frame is "
        "labelled Safe, since a risk assessment requires both objects. Otherwise, for "
        "every child–hazard pair the Euclidean distance between box centroids is "
        "computed and normalised by the image diagonal, d = ‖c_child − c_hazard‖ / "
        "√(W²+H²). The minimum d over all pairs is compared against a calibrated "
        "threshold: below it the frame is Unsafe. Normalising by the diagonal makes d "
        "dimensionless and comparable across resolutions, which a raw pixel threshold "
        "is not.")
    para(doc, "**Design justifications.** The child class and the twelve hazard "
        "classes differ greatly in scale, appearance and frequency; training them "
        "jointly forces one backbone to allocate capacity across both and creates a "
        "class-imbalance problem between one common class and twelve rarer ones. "
        "Training separately removes that coupling and allows independent "
        "configuration — which mattered, since the detectors converged on different "
        "optimal resolutions. Keeping fusion parameter-free means the geometric stage "
        "can be evaluated in isolation: any failure is attributable to the detectors "
        "or the rule, not to a third trained component.")
    para(doc, "**Training procedure.** Both detectors are fine-tuned from "
        "COCO-pretrained YOLOv8-nano weights for 100 epochs at batch size 16 with "
        "default augmentation, using AdamW and the standard YOLOv8 composite loss "
        "(box regression, classification, distribution focal loss). The child detector "
        "uses an initial learning rate of 2×10⁻³ and the hazard detector 6.25×10⁻⁴, "
        "both annealed to 1% of the initial rate over the scheduled epoch count.")
    para(doc, "These rates come from the framework’s own heuristic for a dataset’s "
        "class count. This is worth stating because the framework applies it "
        "*silently*: with the optimizer left at its automatic setting, any "
        "user-specified learning rate is discarded and replaced, and the run’s "
        "configuration file records the supplied value rather than the one used. This "
        "invalidates any hyperparameter search conducted under that setting, so all "
        "tuning reported here names the optimizer explicitly.")

    # ---- 4 Experimental Setup
    heading(doc, "4 Experimental Setup")
    para(doc, "**Datasets.** Two public Roboflow Universe datasets, each used with its "
        "creator’s predefined split without reshuffling (Table 1).")
    table(doc, ["", "Child", "Hazard"], [
        ["Images", "4,705", "5,917"],
        ["Unique source photos", "~649 (≈7 copies)", "3,863 (≈1.5)"],
        ["Classes", "1", "12"],
        ["Train / val / test", "4,080 / 372 / 253", "4,758 / 773 / 386"],
        ["Distributed augmentation", "salt-and-pepper 5%", "50% h-flip"],
    ], "Table 1: Dataset composition.")
    para(doc, "**Two dataset defects were found and are reported rather than worked "
        "around.** Filenames in the child dataset encode the source photograph, making "
        "duplication checkable: 386 source names span the training and evaluation "
        "boundary, and pixel comparison confirms ~17% of validation and ~21% of test "
        "are near-duplicates of training images. The effect was measured, not assumed "
        "— re-evaluating the same weights on a decontaminated validation subset (313 of "
        "372) gives 0.9465 mAP50 against 0.9548 on the full split, an inflation of "
        "0.008. The contamination is real but small, and the dataset was retained. The "
        "hazard dataset is free of cross-split duplication. Separately, the 4,705 "
        "child images derive from only ~649 unique photographs, so effective visual "
        "diversity is an order of magnitude below the image count.")
    para(doc, "Hazard classes are heavily imbalanced: validation support ranges from "
        "367 instances (Coin) to 5 (Chisel). Classes below about 30 instances produce "
        "per-class figures that move substantially on single detections.")
    para(doc, "**Risk-classification evaluation set.** Neither source dataset contains "
        "frames with a child and a hazard together, so neither can evaluate fusion. A "
        "purpose-built set of 200 composited images was constructed by pasting hazard "
        "crops onto held-out child photographs at controlled separations (123 "
        "validation / 77 test, split by source photograph). Ground truth uses a "
        "*reachability* criterion — whether the gap between boxes is at most half the "
        "child’s box height — deliberately chosen because it is not a rescaling of the "
        "distance measure under test, so the evaluation is not circular.")
    para(doc, "**Baselines.** The primary baseline is internal: both detectors trained "
        "at framework-default hyperparameters with the automatic optimizer for the same "
        "100 epochs. This determines whether tuning helped. Four published systems are "
        "used as context; none is directly comparable, and Table 5 records why.")
    para(doc, "**Evaluation metrics.** Average precision is the area under the "
        "precision–recall curve for a class; mAP is its unweighted mean across classes. "
        "mAP50 uses a fixed intersection-over-union threshold of 0.5; mAP50-95 averages "
        "over thresholds 0.5 to 0.95 in steps of 0.05 and is therefore more sensitive "
        "to localisation precision. Both are reported: mAP50 indicates whether the "
        "object is found, mAP50-95 whether it is found accurately. Because mAP is "
        "unweighted, a rare class influences the hazard figure as much as a common one.")
    para(doc, "For risk classification we report **balanced accuracy**, the mean of the "
        "two class recalls, BA = ½(TP/(TP+FN) + TN/(TN+FP)). This is used in preference "
        "to plain accuracy because the evaluation set is ~74% Unsafe, so a degenerate "
        "always-Unsafe classifier scores 0.740 while contributing nothing; balanced "
        "accuracy scores any constant classifier 0.500. Precision, recall and "
        "specificity are reported alongside, since a missed hazard and an unnecessary "
        "alert carry very different costs. Latency is reported per detector and summed, "
        "since both run on every frame.")
    para(doc, "**Implementation.** Ultralytics YOLOv8 pinned at 8.4.106; YOLOv8-nano "
        "(3.01M parameters, 8.1 GFLOPs); NVIDIA Tesla T4 on Google Colab; 100 epochs at "
        "batch 16; AdamW; default loss weights (box 7.5, cls 0.5, dfl 1.5); child at "
        "416 and hazard at 640. Training took 86 and 145 minutes respectively. The "
        "version is pinned because minor versions differ in detection-head "
        "initialisation, shifting the same weights’ mAP50 by about 0.004 — small, but "
        "enough to corrupt an ablation.")
    para(doc, "**Reproducibility.** All code, configurations and result files are in "
        "the project repository. Both datasets are public and fetched at pinned version "
        "numbers so a later re-release cannot silently change the split; two "
        "post-download fixes are scripted (the distributed configuration files contain "
        "broken relative paths, and the child dataset ships as two classes where the "
        "second is mislabelled instances of the first). All runs use seed 0 with "
        "deterministic mode; two independently launched runs of one configuration "
        "reproduced each other to five decimal places. The final training configuration "
        "is embedded in the training script rather than passed on the command line, so "
        "a final run cannot silently disagree with the ablations that chose it. Every "
        "table and figure is regenerated from the result files by a single script.")

    # ---- 5 Results
    heading(doc, "5 Results")
    para(doc, "All development used validation splits. The test splits were used "
        "exactly once, at the end, with every parameter fixed in advance.")
    para(doc, "**Main results.**")
    table(doc, ["System", "Split", "mAP50", "mAP50-95"], [
        ["Child detector", "child test", "0.9670", "0.9081"],
        ["Hazard detector", "hazard test", "0.5506", "0.3958"],
    ], "Table 2: Detector performance on held-out test data.")
    table(doc, ["Risk classification (test)", "Value"], [
        ["Balanced accuracy", "0.5947"],
        ["95% confidence interval", "[0.483, 0.707]"],
        ["Plain accuracy", "0.5974"],
        ["Unsafe precision / recall", "0.683 / 0.609"],
        ["Safe recall", "0.581"],
        ["Confusion TP/FP/FN/TN", "28 / 13 / 18 / 18"],
    ], "Table 3: Risk classification. The interval includes chance (0.500).")
    para(doc, "The child detector’s test figure exceeds its validation figure (0.967 "
        "against 0.955). This is most plausibly contamination rather than "
        "generalisation: the child test split carries the highest measured duplication "
        "(≈21%, against ≈17% in validation).")
    para(doc, "**Computational cost.** Child 3.05 ms and hazard 4.19 ms per image on a "
        "T4, giving **7.24 ms per frame** for the complete pipeline, about 138 frames "
        "per second. Running two specialist detectors instead of one is therefore not a "
        "meaningful cost objection at this model scale.")

    para(doc, "**Ablation: input resolution.** Each detector was trained at 416, 640 "
        "and 832 with all else fixed (Figure 2).")
    figure(doc, "ablation_imgsz.png",
           "Figure 2: Resolution ablation. Right: the child ranking inverts once the "
           "learning rate is corrected — 416 moves from last to first.")
    table(doc, ["imgsz", "Hazard", "Child (lr 1e-2)", "Child (lr 2e-3)"], [
        ["416", "0.5115", "0.9209", "0.9626"],
        ["640", "0.5259", "0.9338", "0.9547"],
        ["832", "0.5003", "0.9150", "0.9504"],
    ], "Table 4: Resolution ablation, mAP50, 30 epochs per cell.")
    para(doc, "Two findings. First, 832 is worse than 640 for both detectors — by 0.034 "
        "mAP50-95 in each case — while costing about 1.6× more inference time; since the "
        "effect replicates across two independent datasets it is more credible than a "
        "single-model result. Second, **the child ranking inverts depending on the "
        "learning rate.** The child detector was never tuned, so the first ablation ran "
        "at the framework default of 1×10⁻², under which 640 wins. Repeated at the "
        "correct 2×10⁻³, every resolution improves and 416 — previously last — wins on "
        "both metrics while running 1.97× faster. The first ablation was internally "
        "consistent, so its comparison was not invalid; but it was conducted at a "
        "handicapped operating point and produced a confidently wrong answer.")

    para(doc, "**Ablation: optimizer.** Each optimizer is given the learning rate "
        "appropriate to it rather than a single shared value. Holding one rate fixed "
        "does not isolate the optimizer, it measures which optimizer happens to suit "
        "that rate — our own tuning data shows the same optimizer scoring 0.2325 mAP50 "
        "at one rate and 0.5020 at another. AdamW is best or tied-best in every "
        "comparison (hazard 0.5259 against SGD 0.5130 and Adam 0.5181; child 0.9547 "
        "against 0.9509 and 0.9528), though the margins of 0.004–0.018 are small.")
    figure(doc, "ablation_optimizer.png",
           "Figure 3: Optimizer ablation, each optimizer at an appropriate "
           "learning rate.")

    para(doc, "**Ablation: hyperparameter tuning.** An evolutionary search over the "
        "initial learning rate and three loss weights ran for six iterations of 20 "
        "epochs; the winner was retrained at the full 100 epochs against the untuned "
        "baseline.")
    table(doc, ["Hazard configuration", "mAP50", "mAP50-95"], [
        ["Framework defaults (baseline)", "0.5657", "0.4072"],
        ["Tuned (lr 8.8e-4, box 8.14, cls 0.75, dfl 1.06)", "0.5256", "0.3913"],
        ["Difference", "−0.0401", "−0.0159"],
    ], "Table 5: Tuning reduced final accuracy.")
    para(doc, "**Tuning made the detector worse**, on 10 of 12 classes. "
        "Hyperparameters were selected on 20-epoch runs and applied to 100-epoch "
        "training; the framework anneals the learning rate across the *scheduled* epoch "
        "count, so a 20-epoch run is fully annealed at its end while a 100-epoch run is "
        "only one fifth through its decay. A configuration that converges well under a "
        "short schedule is not necessarily best over five times as long. The shipped "
        "hazard detector is therefore the untuned baseline. A second limitation: the "
        "learning rate dominated the search, moving fitness by 0.170 while the entire "
        "spread across the three loss weights was 0.034, so the loss weights were never "
        "resolved.")

    para(doc, "**Ablation: distance reference.** The proposal anticipated that "
        "nearest-edge distance might outperform centroid distance, since a large child "
        "box and a small hazard box can have distant centroids while nearly touching.")
    table(doc, ["Detector confidence", "Centroid", "Nearest-edge", "Difference"], [
        ["0.25 (default)", "0.6683", "0.6425", "centroid +0.026"],
        ["0.05", "0.6705", "0.6726", "edge +0.002"],
        ["0.01", "0.6126", "0.5307", "centroid +0.082"],
    ], "Table 6: Balanced accuracy on validation. The ranking flips with a "
       "detector parameter.")
    para(doc, "**The ranking flips with the detector confidence threshold** — a "
        "parameter of the detectors, not of the fusion layer. No ordering survives that "
        "change, so the two references perform equivalently and the original centroid "
        "design is retained. The same experiment produced a more consequential finding: "
        "at the framework’s default confidence the hazard detector **fails to find the "
        "hazard in 57% of images**, and the fusion rule labels all of those Safe without "
        "computing any distance. Under those conditions both measures score below the "
        "degenerate always-Unsafe baseline on plain accuracy. Lowering confidence to "
        "0.05 reduces the failure rate to 22%. Detector confidence therefore affects "
        "risk-classification accuracy more than the distance measure does, and must be "
        "reported as a system parameter.")
    figure(doc, "distance_distributions.png",
           "Figure 4: Distance distributions by class on validation. The classes "
           "overlap substantially under both measures.")

    para(doc, "**Statistical significance.** This is the weakest aspect of the "
        "evaluation and is stated plainly. The risk-classification result carries a 95% "
        "confidence interval of [0.483, 0.707] around a balanced accuracy of 0.5947, "
        "from the normal approximation to the two class recalls on 77 test images. "
        "**The interval includes chance.** The point estimate is above 0.500 and the "
        "validation result (0.6705) was more clearly so, but the held-out evidence does "
        "not establish that the geometric stage performs better than chance. For the "
        "detector results, **no variance estimate exists**: all runs use a fixed seed "
        "with deterministic execution and no configuration was run more than once. "
        "Differences of 0.01–0.02 mAP in the ablations therefore cannot be "
        "distinguished from seed noise, and comparisons at that magnitude — notably the "
        "optimizer ranking — are indicative only. Repeating one configuration across "
        "three seeds would establish an error bar and is the single most valuable "
        "addition to this evaluation.")
    figure(doc, "risk_fusion_confusion_test.png",
           "Figure 5: Risk-classification confusion matrix on the held-out test "
           "split.", width_in=2.5)

    para(doc, "**Qualitative analysis.** Per-class hazard performance on test ranges "
        "from 0.995 (Screwdriver) to 0.168 (Chisel). The spread tracks annotation "
        "support rather than intrinsic difficulty: both extremes rest on a handful of "
        "instances and are equally unreliable. Among well-supported classes, Coin "
        "(0.748), Knife (0.761) and Fork (0.697) perform consistently, while Drink "
        "(0.289) underperforms its support — plausibly because the class groups bottles, "
        "cups and glasses under one label.")
    figure(doc, "per_class_hazard_test.png",
           "Figure 6: Per-class hazard detection on test. The spread tracks "
           "annotation support, not difficulty.")
    para(doc, "The dominant failure mode is a **missed small hazard**, which the fusion "
        "rule converts directly into a false negative: 28 of 77 test frames had no "
        "hazard detected and were labelled Safe without any distance being computed, "
        "accounting for most of the 18 false negatives. In a safety application this is "
        "the more harmful error direction. A second failure mode motivated the "
        "distance-reference ablation: where a child is bent over a small object held in "
        "their hands, the child’s box fills much of the frame, so centroid distance "
        "reads as roughly 0.4 — apparently distant — while the boxes nearly touch. "
        "Nearest-edge distance handles these correctly; that the two measures perform "
        "equivalently overall indicates such cases are outnumbered by ones where "
        "detection error dominates.")
    figure(doc, "final_child_qualitative_pred.jpg",
           "Figure 7: Child detector predictions on validation images.")
    figure(doc, "baseline_hazard_qualitative_pred.jpg",
           "Figure 8: Hazard detector predictions. Small objects at low confidence "
           "are the dominant failure mode.")

    # ---- 6 Discussion
    heading(doc, "6 Discussion")
    para(doc, "**What the results say.** The question was whether two independently "
        "trained specialists combined by a geometric rule form a workable risk "
        "detector, and at what cost. The answer separates into three parts. The "
        "detection stage works and the dual-model cost is negligible: 0.967 and 0.551 "
        "mAP50, running together in 7.24 ms per frame, so the concern that motivated "
        "the question is not borne out. The geometric stage is not confirmed: balanced "
        "accuracy of 0.595 with an interval spanning chance means the proximity rule "
        "cannot be claimed to work. And the bottleneck is detection, not geometry — the "
        "rule can only act when both objects are found, and the hazard detector misses "
        "more than half of them at default confidence. This reframes the problem: the "
        "limiting factor is small-object detection under realistic conditions, not the "
        "distance formula.")
    para(doc, "**Limitations.** The risk evaluation set is composited: hazard crops "
        "pasted onto child photographs at controlled separations. This gives exact "
        "geometry and non-circular ground truth, but pasted objects have visible "
        "boundaries and no lighting or perspective matching, so they test the "
        "pipeline’s mechanics rather than establishing that proximity predicts "
        "real-world danger. The ground-truth criterion embeds an assumption — “Unsafe” "
        "is a gap of at most half the child’s box height, a stated modelling choice, and "
        "not a constant physical distance since the dataset annotates a head in roughly "
        "40% of images and a full body in 57%. Seventy-seven test frames is too few to "
        "resolve the effect. No variance estimates exist. Hyperparameters, resolution "
        "and optimizer were chosen sequentially, each conditioned on the previous "
        "winner, and the resolution ablation shows that assumption failing in practice. "
        "Both the duplication and the low source diversity in the child dataset bound "
        "how strongly generalisation can be claimed.")
    para(doc, "**Ethical considerations.** This system must not be presented as a "
        "substitute for supervision. Its measured false-negative rate is high — it "
        "missed 18 of 46 unsafe frames — and a caregiver who trusts it and reduces "
        "direct attention would be worse off than one who does not use it. Any "
        "deployment must frame the output as a supplementary alert and disclose the "
        "false-negative rate in terms users can act on.")
    para(doc, "The system also requires continuous camera coverage of spaces where "
        "children are, generating exactly the footage that is most sensitive if "
        "breached and creating a record of a child’s movements they cannot consent to. "
        "On-device processing without retention would mitigate this; the present work "
        "does not implement it.")
    para(doc, "**Demographic bias is unmeasured and likely present.** The child dataset "
        "derives from ~649 web-sourced photographs with no recorded distribution of "
        "skin tone, age, body size, clothing or setting. A detector trained on such a "
        "set may perform unevenly across groups, and in this application uneven "
        "performance means uneven protection — a failure falling hardest on children "
        "least represented in the data. We did not measure this, and that omission is "
        "itself a limitation; a deployable version would require a demographically "
        "documented evaluation set.")
    para(doc, "Hazard definition is also cultural and contextual: Coin is a hazard for "
        "a toddler and not for an eight-year-old, and Drink may be a glass of water. A "
        "system that alerts indiscriminately produces false alarms that erode trust, "
        "and alert fatigue in a safety system is itself a safety problem. Finally, "
        "continuous child-tracking infrastructure can be repurposed for monitoring "
        "domestic workers or family members, and the same detect-and-measure pattern "
        "generalises to surveillance unrelated to safety.")

    # ---- 7 Conclusion
    heading(doc, "7 Conclusion")
    para(doc, "We built and evaluated a child safety risk detector from two "
        "independently trained nano-scale detectors and a parameter-free geometric "
        "fusion rule. The detectors perform competitively — 0.967 and 0.551 mAP50 on "
        "held-out test data — and run together in 7.24 ms per frame, showing that the "
        "dual-specialist design carries no meaningful computational penalty at this "
        "scale. The risk-classification stage reached 0.595 balanced accuracy with a "
        "confidence interval that includes chance, so the geometric rule is not "
        "confirmed to work.")
    para(doc, "The most useful finding is diagnostic rather than architectural. "
        "Detection failure, not distance measurement, dominates the error: the hazard "
        "detector misses more than half the hazards at default confidence, and each "
        "miss becomes a false negative before geometry is consulted. The distance "
        "reference — the design question the project set out to examine — proved not to "
        "matter, with the ranking between centroid and edge measures flipping according "
        "to a detector parameter.")
    para(doc, "Three methodological results carry beyond this application. "
        "Hyperparameters selected on a shortened schedule transferred poorly and "
        "*reduced* final accuracy by 0.040 mAP50, because the learning-rate schedule is "
        "defined over the scheduled epoch count. A resolution ablation inverted its "
        "conclusion once an earlier hyperparameter was corrected, showing that "
        "sequential ablation can give confidently wrong answers. And a widely used "
        "framework silently discards a user-specified learning rate under its automatic "
        "optimizer setting, invalidating any search run that way.")
    para(doc, "**Future work**, in order of expected value:")
    bullets(doc, [
        "**Improve small-hazard detection**, the binding constraint. A larger "
        "backbone, higher resolution for the hazard branch alone, or a detector "
        "designed for small objects would address the dominant error directly.",
        "**Build a real, human-labelled co-occurrence evaluation set.** Only this can "
        "establish that proximity predicts genuine danger rather than constructed "
        "geometry, and only this can settle the distance-reference question "
        "non-circularly.",
        "**Report variance** by repeating configurations across seeds, so the small "
        "differences in the ablations can be interpreted at all.",
        "**Treat detector confidence as a tuned parameter**, since it affects "
        "risk-classification accuracy more than the fusion design does.",
        "**Replace the binary label with a graded risk score**, letting the system "
        "express uncertainty rather than committing to Safe on a missed detection — "
        "the failure mode that currently causes most harm.",
    ])

    # ---- back matter
    caps_heading(doc, "ACKNOWLEDGMENTS")
    para(doc, "Insert paragraph text here.", size=SMALL_PT)
    caps_heading(doc, "REFERENCES")
    for r in REFERENCES:
        p = para(doc, r, size=REF_PT, space_after=2)
        p.paragraph_format.left_indent = Inches(0.2)
        p.paragraph_format.first_line_indent = Inches(-0.2)


REFERENCES = [
    "[1]\tMuhammad Haseeb Ahmad, Uzair Ud Din, Fiaz Hussain, Muhammad Farooq, "
    "Ahmad Khan, and Ihtisham Ul Haq. 2025. A Computer Vision Based Child Safety "
    "Solution Using YOLOv8 Architecture. International Journal of Innovations in "
    "Science & Technology 7, 7 (2025), 297–306.",
    "[2]\tAhmed K. AlMhdawi, Nonso Nnamoko, and Ahmed M. Ubaid. 2026. Intelligent "
    "Spatial Estimation for Fire Hazards in Engineering Sites: An Enhanced "
    "YOLOv8-Powered Proximity Analysis Framework. arXiv:2603.09069.",
    "[3]\tRoman Solovyev, Weimin Wang, and Tatiana Gabruseva. 2021. Weighted boxes "
    "fusion: Ensembling boxes from different object detection models. Image and "
    "Vision Computing 107 (2021), 104117.",
    "[4]\tNur Ramadan, Muhammad Muhtadi, Muhammad Z. Rafi, and Retno E. Waluyo. "
    "2025. Risk Detection System for Children Putting Objects into Mouth Based on "
    "Computer Vision using YOLOv11n. Indonesian Journal of Innovation Studies 26, 4 "
    "(2025).",
    "[5]\tFarhan A. Khan and Anup Dey. 2024. Towards enhancing child safety: A deep "
    "learning approach to detect child safe and unsafe objects. In 2024 IEEE "
    "International Women in Engineering (WIE) Conference on Electrical and Computer "
    "Engineering (WIECON-ECE). IEEE, 123–128.",
]


if __name__ == "__main__":
    main()
